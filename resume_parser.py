"""
resume_parser.py
Extracts and cleans text from PDF/DOCX resumes
"""

import re
import io
import PyPDF2
import docx
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt', quiet=True)

try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords', quiet=True)

try:
    nltk.download('punkt_tab', quiet=True)
except:
    pass


# Common sections to identify in resumes
RESUME_SECTIONS = [
    'education', 'experience', 'skills', 'projects',
    'certifications', 'summary', 'objective', 'achievements',
    'work experience', 'technical skills', 'languages', 'interests'
]

# Irrelevant patterns to remove
IRRELEVANT_PATTERNS = [
    r'references available.*',
    r'page \d+ of \d+',
    r'curriculum vitae',
    r'^\s*$',  # empty lines
]


def extract_text_from_pdf(file) -> str:
    """Extract text from a PDF file object."""
    text = ""
    try:
        if hasattr(file, 'read'):
            pdf_reader = PyPDF2.PdfReader(file)
        else:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file))

        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    except Exception as e:
        text = f"Error reading PDF: {str(e)}"
    return text


def extract_text_from_docx(file) -> str:
    """Extract text from a DOCX file object."""
    text = ""
    try:
        if hasattr(file, 'read'):
            doc = docx.Document(file)
        else:
            doc = docx.Document(io.BytesIO(file))

        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        text = f"Error reading DOCX: {str(e)}"
    return text


def extract_text(file, filename: str) -> str:
    """Auto-detect file type and extract text."""
    filename_lower = filename.lower()
    if filename_lower.endswith('.pdf'):
        return extract_text_from_pdf(file)
    elif filename_lower.endswith('.docx') or filename_lower.endswith('.doc'):
        return extract_text_from_docx(file)
    else:
        # Try reading as plain text
        try:
            if hasattr(file, 'read'):
                return file.read().decode('utf-8', errors='ignore')
            return str(file)
        except:
            return ""


def clean_text(text: str) -> str:
    """Clean and normalize extracted text."""
    # Lowercase
    text = text.lower()

    # Remove irrelevant patterns
    for pattern in IRRELEVANT_PATTERNS:
        text = re.sub(pattern, '', text, flags=re.IGNORECASE | re.MULTILINE)

    # Remove special characters but keep letters, numbers, spaces
    text = re.sub(r'[^a-zA-Z0-9\s\+\#\.]', ' ', text)

    # Collapse multiple spaces/newlines
    text = re.sub(r'\s+', ' ', text).strip()

    return text


def extract_email(text: str) -> str:
    """Extract email address from text."""
    pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    match = re.search(pattern, text)
    return match.group(0) if match else "Not found"


def extract_phone(text: str) -> str:
    """Extract phone number from text."""
    pattern = r'(\+?\d[\d\s\-\(\)]{8,15}\d)'
    match = re.search(pattern, text)
    return match.group(0).strip() if match else "Not found"


def extract_name(text: str) -> str:
    """Extract candidate name (first non-empty line heuristic)."""
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    if lines:
        # First line is usually the name
        name = lines[0]
        # Filter out if it looks like an email or URL
        if '@' not in name and 'http' not in name and len(name.split()) <= 5:
            return name.title()
    return "Unknown"


def extract_skills_from_text(text: str, skill_keywords: list) -> list:
    """Extract skills present in resume text."""
    text_lower = text.lower()
    found_skills = []
    for skill in skill_keywords:
        if skill.lower() in text_lower:
            found_skills.append(skill)
    return found_skills


def parse_resume(file, filename: str) -> dict:
    """
    Full pipeline: extract → clean → parse metadata.
    Returns a dict with raw_text, clean_text, name, email, phone.
    """
    raw_text = extract_text(file, filename)
    cleaned = clean_text(raw_text)

    return {
        'filename': filename,
        'raw_text': raw_text,
        'clean_text': cleaned,
        'name': extract_name(raw_text),
        'email': extract_email(raw_text),
        'phone': extract_phone(raw_text),
        'char_count': len(cleaned),
        'word_count': len(cleaned.split())
    }
